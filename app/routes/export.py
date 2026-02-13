import io
from flask import Blueprint, send_file
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from app.db import get_db

bp = Blueprint('export', __name__)


def set_rtl_paragraph(paragraph):
    pPr = paragraph._element.get_or_add_pPr()
    bidi = pPr.makeelement(qn('w:bidi'), {})
    pPr.append(bidi)


def set_rtl_run(run, font_name='David', font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    rPr = run._element.get_or_add_rPr()
    rtl_elem = rPr.makeelement(qn('w:rtl'), {})
    rPr.append(rtl_elem)


@bp.route('/exam/<int:exam_id>/docx')
def export_docx(exam_id):
    db = get_db()
    exam = db.execute("SELECT * FROM exams WHERE id=?", (exam_id,)).fetchone()
    questions = db.execute("""
        SELECT eq.position, q.*
        FROM exam_questions eq
        JOIN questions q ON q.id = eq.question_id
        WHERE eq.exam_id = ?
        ORDER BY eq.position
    """, (exam_id,)).fetchall()
    doc = Document()

    # Document-level RTL
    sect = doc.sections[0]._sectPr
    bidi = sect.makeelement(qn('w:bidi'), {})
    sect.append(bidi)

    # Title
    title_p = doc.add_paragraph()
    set_rtl_paragraph(title_p)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(exam['title'])
    set_rtl_run(run, font_size=18)
    run.bold = True

    if exam['description']:
        desc_p = doc.add_paragraph()
        set_rtl_paragraph(desc_p)
        desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = desc_p.add_run(exam['description'])
        set_rtl_run(run, font_size=11)

    doc.add_paragraph()  # spacing

    # Questions
    for q in questions:
        pos = q['position']
        q_para = doc.add_paragraph()
        set_rtl_paragraph(q_para)
        run = q_para.add_run(f'{pos}. {q["stem_he"]}')
        set_rtl_run(run)
        run.bold = True

        for letter, field in [('א', 'option_a'), ('ב', 'option_b'),
                               ('ג', 'option_c'), ('ד', 'option_d')]:
            text = q[field]
            if text:
                opt_p = doc.add_paragraph()
                set_rtl_paragraph(opt_p)
                opt_p.paragraph_format.left_indent = Cm(1)
                run = opt_p.add_run(f'   {letter}. {text}')
                set_rtl_run(run)

        if q['option_e']:
            opt_p = doc.add_paragraph()
            set_rtl_paragraph(opt_p)
            opt_p.paragraph_format.left_indent = Cm(1)
            run = opt_p.add_run(f'   ה. {q["option_e"]}')
            set_rtl_run(run)

        doc.add_paragraph()

    # Answer key
    doc.add_page_break()
    key_p = doc.add_paragraph()
    set_rtl_paragraph(key_p)
    key_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = key_p.add_run('מפתח תשובות')
    set_rtl_run(run, font_size=16)
    run.bold = True

    doc.add_paragraph()

    answer_map = {'A': 'א', 'B': 'ב', 'C': 'ג', 'D': 'ד', 'E': 'ה'}

    for q in questions:
        ans_p = doc.add_paragraph()
        set_rtl_paragraph(ans_p)
        correct_he = answer_map.get(q['correct_answer'], q['correct_answer'])
        run = ans_p.add_run(f'{q["position"]}. {correct_he}')
        set_rtl_run(run)
        run.bold = True

        if q['explanation_he']:
            exp_p = doc.add_paragraph()
            set_rtl_paragraph(exp_p)
            exp_p.paragraph_format.left_indent = Cm(0.5)
            run = exp_p.add_run(q['explanation_he'])
            set_rtl_run(run, font_size=10)

        doc.add_paragraph()

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f'{exam["title"]}.docx'
    return send_file(buffer, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
